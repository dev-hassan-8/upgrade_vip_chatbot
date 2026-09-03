from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    source: str
    text: str
    document_type: str
    metadata: dict


def discover_documents(directory: Path) -> list[Path]:
    if not directory.exists():
        return []
    return sorted(
        path
        for path in directory.iterdir()
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return _load_txt(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    raise ValueError(f"Unsupported file format: {path.suffix}")


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text.strip()


def _load_txt(path: Path) -> LoadedDocument:
    text = clean_text(path.read_text(encoding="utf-8"))
    return LoadedDocument(
        source=path.name,
        text=text,
        document_type="txt",
        metadata={"filename": path.name},
    )


def _load_pdf(path: Path) -> LoadedDocument:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {index}]\n{page_text}")

    text = clean_text("\n\n".join(pages))
    return LoadedDocument(
        source=path.name,
        text=text,
        document_type="pdf",
        metadata={"filename": path.name, "page_count": len(reader.pages)},
    )


def _load_docx(path: Path) -> LoadedDocument:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    text = clean_text("\n\n".join(paragraphs))
    return LoadedDocument(
        source=path.name,
        text=text,
        document_type="docx",
        metadata={"filename": path.name},
    )
